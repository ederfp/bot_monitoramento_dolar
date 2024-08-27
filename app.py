from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.common.exceptions import *
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions

from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import win32com.client
from datetime import datetime
from time import sleep
import os


class BotMonitoringDollar:

    def __init__(self) -> None:
        self.date_now = datetime.now().strftime('%d/%m/%Y')
        self.text_quote_dollar = ''
        self.site_quote = 'https://www.confidencecambio.com.br/'
        self.path = os.path.dirname(os.path.abspath(__file__))
        self.username = os.environ['USERNAME']
        self.path_download = os.path.join(r'C:\Users', self.username, 'Downloads', 'Arquivos Bot_Dolar')
        
        os.makedirs(self.path_download, exist_ok=True)

    def start_driver(self):
        chrome_options = Options()

        arguments = ['--lang=pt-BR', '--window-size=800,600',
                    '--incognito']

        for argument in arguments:
            chrome_options.add_argument(argument)

        #caminho_padrao_para_download = 'E:\\Storage\\Desktop'

        chrome_options.add_experimental_option("prefs", {
            #'download.default_directory': caminho_padrao_para_download,
            #'download.directory_upgrade': True,
            'download.prompt_for_download': False,
            "profile.default_content_setting_values.notifications": 2,
            "profile.default_content_setting_values.automatic_downloads": 1,
        })

        driver = webdriver.Chrome(options=chrome_options)

        wait = WebDriverWait(
            driver,
            10,
            poll_frequency=1,
            ignored_exceptions=[
                NoSuchElementException,
                ElementNotVisibleException,
                ElementNotSelectableException,
            ]
        )

        return driver, wait

    def get_dollar_info(self):
        driver, wait = BotMonitoringDollar.start_driver(self)
        try:
            driver.get('https://www.confidencecambio.com.br/')
            sleep(5)
            driver.maximize_window()
        except:
            print('Erro para Carregar o Site...Aguardar nova tentativa.')
            BotMonitoringDollar.get_dollar_info(self)

        try:
            cotacao_dolar = wait.until(expected_conditions.visibility_of_any_elements_located((
                    By.XPATH, '//strong[@class="p"]')))
        except:
            print('Valor Dolar não foi encontrado.')
        
        #Print(imagem) do site onde a cotação foi realizada
        caminho_print = os.path.join(self.path_download, 'print site cotação.jpg')
        driver.save_screenshot(caminho_print)
        sleep(2)

        for texto in cotacao_dolar:
            self.text_quote_dollar = texto.text
            break
        
        driver.close()

    def file_word(self):
        documento = Document()

        title = documento.add_heading(f'Cotação Atual do Dólar - {self.text_quote_dollar} ({self.date_now})', 1)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        documento.add_paragraph("")
        paragrafo1 = documento.add_paragraph(f'O dólar está no valor de ')
        paragrafo1.add_run(f'{self.text_quote_dollar}').bold = True
        paragrafo1.add_run(', na data ')
        paragrafo1.add_run(f'{self.date_now}.').bold = True

        paragrafo2 = documento.add_paragraph(f'Valor cotado no site ')
        paragrafo2.add_run(f'{self.site_quote}.').bold = True

        documento.add_paragraph('Print da cotação atual.')

        caminho_picture = os.path.join(self.path_download, 'print site cotação.jpg')
        documento.add_picture(caminho_picture, width=Cm(15))

        paragrafo3 = documento.add_paragraph('Cotação feita por - ')
        paragrafo3.add_run('Eder Fornielles').italic = True

        path_save_doc = os.path.join(self.path_download, 'cotação dolar.docx')
        documento.save(path_save_doc)

    def file_pdf(self):
        wdFormatPDF = 17

        entrada = os.path.join(self.path_download, 'cotação dolar.docx')
        saida = os.path.join(self.path_download, 'cotação dolar pdf.pdf')

        word = win32com.client.Dispatch('Word.Application')
        doc = word.Documents.Open(entrada)
        doc.SaveAs(saida, FileFormat=wdFormatPDF)
        doc.Close()
        word.Quit()


self = BotMonitoringDollar()
self.get_dollar_info()
self.file_word()
self.file_pdf()